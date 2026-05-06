"""DOCX → JSON parser using python-docx."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from .base_parser import BaseParser


_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}
_BOLD_HEADING_RE = re.compile(r"^(\d+(?:\.\d+)*\.?)\s+(.+)$")

# XML tag names used to distinguish paragraphs from tables in body order
_TAG_PARA        = qn("w:p")
_TAG_TABLE       = qn("w:tbl")
_TAG_SDT         = qn("w:sdt")          # structured document tag (content control)
_TAG_SDT_CONTENT = qn("w:sdtContent")
_TAG_RUN_TEXT    = qn("w:t")


def _table_to_text(table: Any) -> str:
    """Convert a python-docx Table to a readable text block."""
    lines: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        # Deduplicate merged cells (python-docx repeats merged cell text)
        deduped: list[str] = []
        prev = object()
        for c in cells:
            if c != prev:
                deduped.append(c)
            prev = c
        line = " | ".join(deduped)
        if line.strip():
            lines.append(line)
    return "\n".join(lines)


def _sdt_text(sdt_elem: Any) -> str:
    """Extract plain text from a w:sdt content-control element."""
    for child in sdt_elem:
        if child.tag == _TAG_SDT_CONTENT:
            return "".join(el.text or "" for el in child.iter(_TAG_RUN_TEXT)).strip()
    return ""


def _iter_body_blocks(doc: Document):
    """
    Yield (kind, obj) in document order from the body XML.
    kind is 'para', 'table', or 'text' (raw string from a content control).
    """
    para_map  = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in doc.element.body:
        tag = child.tag
        if tag == _TAG_PARA and child in para_map:
            yield "para", para_map[child]
        elif tag == _TAG_TABLE and child in table_map:
            yield "table", table_map[child]
        elif tag == _TAG_SDT:
            # Paragraphs inside content controls are not in doc.paragraphs;
            # extract raw text so nothing is silently dropped.
            text = _sdt_text(child)
            if text:
                yield "text", text


class DocxParser(BaseParser):
    def __init__(self, file_path: str | Path) -> None:
        super().__init__(file_path)
        self._doc: Document | None = None

    def _load(self) -> Document:
        if self._doc is None:
            import zipfile
            try:
                self._doc = Document(str(self.file_path))
            except Exception as exc:
                # Check if the file is a legacy .doc binary (starts with D0 CF magic bytes)
                try:
                    magic = self.file_path.read_bytes()[:4]
                except OSError:
                    magic = b""
                if magic[:2] == b"\xd0\xcf":
                    raise ValueError(
                        "This appears to be a legacy .doc (Word 97) file. "
                        "Please convert it to .docx format and re-upload."
                    ) from exc
                raise ValueError(f"Could not open DOCX file: {exc}") from exc
        return self._doc

    # ── raw text ───────────────────────────────────────────────────────────────

    def extract_raw_text(self) -> str:
        doc = self._load()
        parts: list[str] = []
        for kind, obj in _iter_body_blocks(doc):
            if kind == "para":
                text = obj.text.strip()
                if text:
                    parts.append(text)
            elif kind == "table":
                table_text = _table_to_text(obj)
                if table_text:
                    parts.append(table_text)
            else:  # text — raw string from a content control
                parts.append(obj)
        return "\n".join(parts)

    # ── sections ──────────────────────────────────────────────────────────────

    def extract_sections(self) -> list[dict[str, Any]]:
        doc = self._load()
        sections: list[dict[str, Any]] = []
        current: dict[str, Any] | None = None
        order = 0

        for kind, obj in _iter_body_blocks(doc):
            if kind == "para":
                text = obj.text.strip()
                if not text:
                    continue

                is_heading = obj.style.name in _HEADING_STYLES or self._is_bold_heading(obj)

                if is_heading:
                    if current:
                        current["content"] = current["content"].strip()
                        sections.append(current)
                    order += 1
                    num_match = _BOLD_HEADING_RE.match(text)
                    current = {
                        "section_id": num_match.group(1).rstrip(".") if num_match else f"s{order}",
                        "title":      num_match.group(2) if num_match else text,
                        "content":    "",
                        "order":      order,
                        "subsections": [],
                    }
                elif current is not None:
                    current["content"] += text + "\n"

            elif kind == "text":  # raw text from a content control
                if obj and current is not None:
                    current["content"] += obj + "\n"

            else:  # table — append its text to the current section
                table_text = _table_to_text(obj)
                if not table_text:
                    continue
                if current is not None:
                    current["content"] += table_text + "\n"
                else:
                    # Table appears before any heading — create an implicit section
                    order += 1
                    current = {
                        "section_id": f"tbl{order}",
                        "title":      "Table",
                        "content":    table_text + "\n",
                        "order":      order,
                        "subsections": [],
                    }

        if current:
            current["content"] = current["content"].strip()
            sections.append(current)

        if not sections:
            raw = self.extract_raw_text()
            sections = self._fallback_sections(raw)

        return sections

    # ── helpers ────────────────────────────────────────────────────────────────

    @staticmethod
    def _is_bold_heading(para: Any) -> bool:
        """Return True if all non-empty runs are bold and the text is short enough to be a heading."""
        runs = [r for r in para.runs if r.text.strip()]
        if not runs or not all(r.bold for r in runs):
            return False
        # Long bold sentences are content, not headings (e.g. "Annually, not to exceed 15 months:")
        text = "".join(r.text for r in runs).strip()
        return len(text.split()) <= 8

    @staticmethod
    def _fallback_sections(raw: str) -> list[dict[str, Any]]:
        paragraphs = [p.strip() for p in re.split(r"\n{2,}", raw) if p.strip()]
        sections: list[dict[str, Any]] = []
        for i, para in enumerate(paragraphs[:50], 1):
            first_line = para.splitlines()[0][:80]
            sections.append(
                {
                    "section_id": f"p{i}",
                    "title":      first_line,
                    "content":    para,
                    "order":      i,
                    "subsections": [],
                }
            )
        return sections
